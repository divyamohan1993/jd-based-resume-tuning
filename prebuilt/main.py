# ==============================================================================  
#                                CONFIDENTIAL  
# ==============================================================================  
#  
#  Title       : SURAT - Shoolini University Resume Analysis & Tuning System
#  Description : A Shoolini University Resume Analysis & Tuning System  
#  Author      : Divya Mohan  
#  Created     : May 26, 2025  
#  Version     : 1.0.0  
#  
#  © 2025 Divya Mohan. ALL RIGHTS RESERVED.  
#  
#  NOTICE TO RECIPIENT:  
#  --------------------  
#  This software and its accompanying documentation (the “Software”) are the  
#  proprietary and confidential property of Divya Mohan. Unauthorized copying,  
#  adaptation, distribution, use, or disclosure of this Software, in whole or in  
#  part, is strictly prohibited. No license or right to the Software is granted  
#  to you by implication, estoppel, or otherwise. Any permitted use must be  
#  pursuant to a written license agreement signed by Divya Mohan.  
#  
#  LEGAL WARNING:  
#  --------------  
#  • You agree that any breach of the terms set forth herein will cause  
#    irreparable harm to the owner, for which monetary damages may be inadequate.  
#  • Divya Mohan reserves the right to seek injunctive relief, damages, and any  
#    other remedies available at law or in equity against any party breaching  
#    these terms.  
#  • This Software is protected under the copyright laws of India,  
#    international treaties, and other intellectual property laws.  
#  • Civil and criminal penalties may apply for unauthorized use or distribution.  
#  
#  DISCLAIMER OF WARRANTY & LIMITATION OF LIABILITY:  
#  ------------------------------------------------  
#  THIS SOFTWARE IS PROVIDED “AS IS,” WITHOUT WARRANTY OF ANY KIND, EXPRESS OR  
#  IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,  
#  FITNESS FOR A PARTICULAR PURPOSE, OR NON-INFRINGEMENT. IN NO EVENT SHALL  
#  DIVYA MOHAN BE LIABLE FOR ANY CLAIM, DAMAGES, OR OTHER LIABILITY, WHETHER IN  
#  AN ACTION OF CONTRACT, TORT, OR OTHERWISE, ARISING FROM, OUT OF, OR IN  
#  CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE SOFTWARE.  
#  
#  OVERVIEW:  
#  ---------  
#  This Flask-based system provides:  
#    1. Input Sanitization:  
#       • Truncates input to safe length, removes control characters, HTML-escapes,  
#         and neutralizes sequences that could break prompts or code fences.  
#    2. Resume Text Extraction:  
#       • Parses PDF and DOCX files via python-magic, PyPDF2, and Mammoth.  
#       • Normalizes bullets, headings, and soft-wrapped lines.  
#    3. Skill Extraction & Categorization:  
#       • Uses Google Gemini Generative AI to extract Technical, Soft, and  
#         Domain skills from job descriptions, formatted as JSON.  
#    4. Resume Analysis:  
#       • Matches candidate skills to required skills, computing match percentages,  
#         emotion-based gap analysis, and category-wise breakdowns.  
#       • Leverages AI for deep contextual matching beyond simple substring checks.  
#    5. Resume Tailoring:  
#       • AI-driven rewriting to highlight relevant skills, achievements, and  
#         keywords in a concise, one-page format.  
#       • Offers PDF/DOCX output in “classic,” “modern,” or “minimal” templates.  
#    6. Export & Delivery:  
#       • Endpoints to download tailored resumes or preview rewritten content.  
#  
#  END-USER NOTES:  
#  ---------------  
#  • Ensure your environment variable GOOGLE_API_KEY is set before running.  
#  • Install dependencies via:  
#       pip install flask flask-cors google-generativeai python-magic PyPDF2 mammoth python-docx reportlab  
#  • To modify or extend functionality, obtain prior written consent from the author.  
#  
# ==============================================================================  

from flask import Flask, request, jsonify, render_template, send_file
import google.generativeai as genai
from flask_cors import CORS
import io
import magic
import json
import mammoth
import PyPDF2
import docx
import itertools
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import re
import html
import os
from dotenv import load_dotenv

# Load variables from .env into os.environ
load_dotenv()

app = Flask(__name__)
CORS(app, origins=["https://jdtuning.dmj.one"])  # Enable CORS for frontend communication

# Now read your key
GEMINI_KEY = os.getenv("GEMINI_KEY")
if not GEMINI_KEY:
    raise RuntimeError("Missing GEMINI_KEY in environment")

# Configure Gemini
genai.configure(api_key=GEMINI_KEY)

geminimodel = "gemini-2.5-pro-preview-05-06"
geminimodel = "gemini-1.5-pro-latest"
geminimodel = "gemini-2.5-flash-preview-05-20"


# Precompile control‐char regex
_CONTROL_CHARS = re.compile(r'[\x00-\x1F\x7F-\x9F]')

def sanitize_input(text: str, max_length: int = 50000) -> str:
    """
    1) Truncate to a safe length.
    2) Strip out C0/C1 control characters.
    3) HTML-escape <, >, &, quotes.
    4) Escape sequences that might break your prompts or code.
    """
    # 1) limit overall length
    text = text[:max_length]

    # 2) drop control characters
    text = _CONTROL_CHARS.sub('', text)

    # 3) HTML-escape so '<script>' becomes '&lt;script&gt;'
    text = html.escape(text)

    # 4) neutralize code‐fences, triple-quotes, backticks, braces, dollars, backslashes
    for seq in ('```', '"""', "'''", '`', '{', '}', '$', '\\'):
        safe = ''.join('\\' + ch for ch in seq)
        text = text.replace(seq, safe)

    return text


def extract_text_from_file(file):
    """
    Extract text from uploaded file (PDF or DOCX) and then normalize spacing & lists.
    """
    def normalize_extracted_text(text: str) -> str:
        """
        Preserve bullets, break on headings (ALL CAPS or ending in ':'),
        join soft-wrapped lines, collapse blanks.
        """
        lines = text.splitlines()
        out = []
        for ln in lines:
            stripped = ln.strip()
            if not stripped:
                # collapse multiple blanks into one
                if out and out[-1] != "":
                    out.append("")
                continue

            # force break before ALL-CAPS headings or lines ending with colon
            is_heading = stripped.isupper() or stripped.endswith(":")
            if is_heading:
                if out and out[-1] != "":
                    out.append("")      # blank line before heading
                out.append(stripped)
                continue

            # list item?
            if stripped.startswith(("-", "*", "•", "–")):
                out.append(stripped)
                continue

            # continuation of previous paragraph?
            if out and out[-1] and not re.search(r"[\.:\?!]$", out[-1]):
                out[-1] = out[-1] + " " + stripped
            else:
                out.append(stripped)

        return "\n".join(out).strip()


    filename = file.filename.lower()
    raw = ""
    # try MIME
    try:
        mime = magic.Magic(mime=True)
        file_mime = mime.from_buffer(file.read(2048))
        file.seek(0)
        if 'pdf' in filename or file_mime == 'application/pdf':
            reader = PyPDF2.PdfReader(file)
            for p in reader.pages:
                raw += p.extract_text() or ""
        elif 'docx' in filename or 'officedocument.wordprocessingml.document' in file_mime:
            try:
                raw = mammoth.convert_to_text(file).value
            except AttributeError:
                raw = mammoth.extract_raw_text(file).value
        else:
            raise ValueError("Unsupported file type.")
    except Exception:
        # fallback on extension
        file.seek(0)
        if filename.endswith('.pdf'):
            reader = PyPDF2.PdfReader(file)
            for p in reader.pages:
                raw += p.extract_text() or ""
        elif filename.endswith('.docx'):
            try:
                raw = mammoth.convert_to_text(file).value
            except AttributeError:
                html = mammoth.convert_to_html(file).value
                text = re.sub(r'<[^>]+>', ' ', html)
                raw = re.sub(r'\s+', ' ', text).strip()
        else:
            raise ValueError("Unsupported file type.")

    # now normalize everything once
    return normalize_extracted_text(raw)     

def convert_to_pdf(text):
    """
    Simple PDF conversion for backward compatibility
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, 
                             rightMargin=72, leftMargin=72, 
                             topMargin=36, bottomMargin=18)  # Reduced top margin
    
    styles = getSampleStyleSheet()
    # Create smaller font size style
    small_style = ParagraphStyle(
        'Small',
        parent=styles['Normal'],
        fontSize=9,  # Smaller font size
        leading=11   # Reduced line spacing
    )
    
    story = []
    
    # Clean text: remove asterisks and excess whitespace
    cleaned_text = text.replace('*', '')
    # Remove multiple consecutive newlines
    import re
    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
    
    # Split text into paragraphs
    paragraphs = cleaned_text.split('\n')
    for para in paragraphs:
        if para.strip():  # Skip empty paragraphs
            p = Paragraph(para, small_style)
            story.append(p)
            # Add smaller spacing between paragraphs
            story.append(Spacer(1, 3))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def convert_to_pdf_classic(text, template_style):
    """
    Convert text to a downloadable PDF with classic template style
    """
    buffer = io.BytesIO()
    # Reduce margins for more content space
    doc = SimpleDocTemplate(buffer, pagesize=letter, 
                             rightMargin=54, leftMargin=54, 
                             topMargin=36, bottomMargin=18)
    
    # Create custom styles based on template with smaller font sizes
    styles = getSampleStyleSheet()
    
    if template_style == "classic":
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Heading1'],
            fontSize=14,  # Smaller title
            spaceAfter=6,  # Reduced spacing
            textColor=colors.darkblue
        )
        heading_style = ParagraphStyle(
            'Heading',
            parent=styles['Heading2'],
            fontSize=12,  # Smaller heading
            spaceBefore=6,  # Reduced spacing
            spaceAfter=3,  # Reduced spacing
            textColor=colors.darkblue
        )
        normal_style = ParagraphStyle(
            'Normal',
            parent=styles['Normal'],
            fontSize=9,  # Smaller text
            leading=11   # Reduced line spacing
        )
    
    elif template_style == "modern":
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Heading1'],
            fontSize=14,  # Smaller title
            spaceAfter=6,  # Reduced spacing
            textColor=colors.darkblue,
            alignment=1  # Center aligned
        )
        heading_style = ParagraphStyle(
            'Heading',
            parent=styles['Heading2'],
            fontSize=12,  # Smaller heading
            spaceBefore=6,  # Reduced spacing
            spaceAfter=3,  # Reduced spacing
            textColor=colors.teal
        )
        normal_style = ParagraphStyle(
            'Normal',
            parent=styles['Normal'],
            fontSize=9,  # Smaller text
            leading=11,  # Reduced line spacing
            textColor=colors.black
        )
    
    elif template_style == "minimal":
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Heading1'],
            fontSize=13,  # Smaller title
            spaceAfter=6,  # Reduced spacing
            textColor=colors.black
        )
        heading_style = ParagraphStyle(
            'Heading',
            parent=styles['Heading2'],
            fontSize=11,  # Smaller heading
            spaceBefore=6,  # Reduced spacing
            spaceAfter=3,  # Reduced spacing
            textColor=colors.gray
        )
        normal_style = ParagraphStyle(
            'Normal',
            parent=styles['Normal'],
            fontSize=9,  # Smaller text
            leading=11   # Reduced line spacing
        )
    
    else:  # Default professional
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Heading1'],
            fontSize=14,  # Smaller title
            spaceAfter=6,  # Reduced spacing
            textColor=colors.black
        )
        heading_style = ParagraphStyle(
            'Heading',
            parent=styles['Heading2'],
            fontSize=12,  # Smaller heading
            spaceBefore=6,  # Reduced spacing
            spaceAfter=3,  # Reduced spacing
            textColor=colors.darkblue
        )
        normal_style = ParagraphStyle(
            'Normal',
            parent=styles['Normal'],
            fontSize=9,  # Smaller text
            leading=11   # Reduced line spacing
        )
    
    story = []
    
    # Clean text: remove asterisks and excess whitespace
    cleaned_text = text.replace('*', '')
    # Remove multiple consecutive newlines
    import re
    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
    
    # Process text - assume sections are separated by double newlines
    sections = cleaned_text.split('\n\n')
    
    # First section is typically contact info/name - treat as title
    if sections:
        story.append(Paragraph(sections[0], title_style))
        story.append(Spacer(1, 6))  # Reduced spacing
    
    # Process remaining sections
    current_heading = None
    current_content = []
    
    for section in sections[1:]:
        # Check if this is a heading (shorter line that ends with a colon or all caps)
        lines = section.split('\n')
        if len(lines) > 0 and (lines[0].isupper() or lines[0].endswith(':') or len(lines[0]) < 30):
            # If we had a previous heading with content, add it to the story
            if current_heading:
                story.append(Paragraph(current_heading, heading_style))
                for content in current_content:
                    if content.strip():  # Skip empty lines
                        story.append(Paragraph(content, normal_style))
                story.append(Spacer(1, 3))  # Reduced spacing
            
            current_heading = lines[0]
            current_content = lines[1:] if len(lines) > 1 else []
        else:
            # This is content for the current heading
            current_content.extend(lines)
    
    # Add the last section
    if current_heading:
        story.append(Paragraph(current_heading, heading_style))
        for content in current_content:
            if content.strip():  # Skip empty lines
                story.append(Paragraph(content, normal_style))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def convert_to_docx(text):
    """
    Simple DOCX conversion for backward compatibility
    """
    document = docx.Document()
    
    # Clean text: remove asterisks and excess whitespace
    cleaned_text = text.replace('*', '')
    # Remove multiple consecutive newlines
    import re
    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
    
    # Make default font smaller
    style = document.styles['Normal']
    style.font.size = Pt(9)
    
    # Set smaller margins for more content space
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Split text into paragraphs
    paragraphs = cleaned_text.split('\n')
    for para in paragraphs:
        if para.strip():  # Skip empty paragraphs
            document.add_paragraph(para)
    
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

def convert_to_docx_template(text, template_style):
    """
    Convert text to a downloadable DOCX with selected template
    """
    document = docx.Document()
    
    # Set smaller margins for all templates to fit more content
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Set up styles based on template
    if template_style == "classic":
        # Set up styles
        document.styles['Normal'].font.name = 'Garamond'
        document.styles['Normal'].font.size = Pt(9)  # Smaller font
        
        # Check if Title style exists before adding
        title_style = document.styles['Title'] if 'Title' in document.styles else document.styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Garamond'
        title_style.font.size = Pt(14)  # Smaller title
        title_style.font.bold = True
        
        # Check if Heading style exists before adding
        heading_style = document.styles['Heading'] if 'Heading' in document.styles else document.styles.add_style('Heading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Garamond'
        heading_style.font.size = Pt(12)  # Smaller heading
        heading_style.font.bold = True
        
    elif template_style == "modern":
        # Set up styles
        document.styles['Normal'].font.name = 'Calibri'
        document.styles['Normal'].font.size = Pt(9)  # Smaller font
        
        # Check if Title style exists before adding
        title_style = document.styles['Title'] if 'Title' in document.styles else document.styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(14)  # Smaller title
        title_style.font.bold = True
        
        # Check if Heading style exists before adding
        heading_style = document.styles['Heading'] if 'Heading' in document.styles else document.styles.add_style('Heading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Calibri'
        heading_style.font.size = Pt(12)  # Smaller heading
        heading_style.font.bold = True
        
    elif template_style == "minimal":
        # Set up styles
        document.styles['Normal'].font.name = 'Arial'
        document.styles['Normal'].font.size = Pt(9)  # Smaller font
        
        # Check if Title style exists before adding
        title_style = document.styles['Title'] if 'Title' in document.styles else document.styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(13)  # Smaller title
        title_style.font.bold = True
        
        # Check if Heading style exists before adding
        heading_style = document.styles['Heading'] if 'Heading' in document.styles else document.styles.add_style('Heading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Arial'
        heading_style.font.size = Pt(11)  # Smaller heading
        heading_style.font.bold = True
        
    else:  # Default professional
        # Set up styles
        document.styles['Normal'].font.name = 'Times New Roman'
        document.styles['Normal'].font.size = Pt(9)  # Smaller font
        
        # Check if Title style exists before adding
        title_style = document.styles['Title'] if 'Title' in document.styles else document.styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Times New Roman'
        title_style.font.size = Pt(14)  # Smaller title
        title_style.font.bold = True
        
        # Check if Heading style exists before adding
        heading_style = document.styles['Heading'] if 'Heading' in document.styles else document.styles.add_style('Heading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Times New Roman'
        heading_style.font.size = Pt(12)  # Smaller heading
        heading_style.font.bold = True
    
    # Clean text: remove asterisks and excess whitespace
    cleaned_text = text.replace('*', '')
    # Remove multiple consecutive newlines
    import re
    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
    
    # Process text - assume sections are separated by double newlines
    sections = cleaned_text.split('\n\n')
    
    # First section is typically contact info/name - treat as title
    if sections:
        title_para = document.add_paragraph(style='Title')
        title_para.add_run(sections[0])
        if template_style == "modern":  # Center align for modern template
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Reduce spacing after title
        title_para.paragraph_format.space_after = Pt(6)
    
    # Process remaining sections
    current_heading = None
    current_content = []
    
    for section in sections[1:]:
        # Check if this is a heading (shorter line that ends with a colon or all caps)
        lines = section.split('\n')
        if len(lines) > 0 and (lines[0].isupper() or lines[0].endswith(':') or len(lines[0]) < 30):
            # If we had a previous heading with content, add it to the document
            if current_heading:
                head_para = document.add_paragraph(style='Heading')
                head_para.add_run(current_heading)
                # Reduce spacing before and after heading
                head_para.paragraph_format.space_before = Pt(6)
                head_para.paragraph_format.space_after = Pt(3)
                
                for content in current_content:
                    if content.strip():  # Skip empty lines
                        para = document.add_paragraph(content, style='Normal')
                        # Reduce spacing between paragraphs
                        para.paragraph_format.space_after = Pt(0)
            
            current_heading = lines[0]
            current_content = lines[1:] if len(lines) > 1 else []
        else:
            # This is content for the current heading
            current_content.extend(lines)
    
    # Add the last section
    if current_heading:
        head_para = document.add_paragraph(style='Heading')
        head_para.add_run(current_heading)
        # Reduce spacing before and after heading
        head_para.paragraph_format.space_before = Pt(6)
        head_para.paragraph_format.space_after = Pt(3)
        
        for content in current_content:
            if content.strip():  # Skip empty lines
                para = document.add_paragraph(content, style='Normal')
                # Reduce spacing between paragraphs
                para.paragraph_format.space_after = Pt(0)
    
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

def extract_skills(job_description):
    """
    Extract skills from job description and categorize them
    """
    model = genai.GenerativeModel(geminimodel)
    prompt = f"""
    Analyze the following job description and extract skills in these categories:
    1. Technical Skills: programming languages, tools, platforms, etc.
    2. Soft Skills: communication, teamwork, leadership, etc.
    3. Domain Knowledge: industry-specific knowledge, regulations, etc.
    
    Format your response as JSON with these three categories as keys and arrays of skills as values.
    
    Job Description:
    {job_description}
    """
    
    response = model.generate_content(prompt)
    skills_text = response.text.strip()
    
    # Process the JSON response
    import json
    try:
        # Try to find JSON content in the response
        json_match = re.search(r'```json(.*?)```', skills_text, re.DOTALL)
        if json_match:
            skills_json = json.loads(json_match.group(1).strip())
        else:
            # If no JSON code block found, try parsing the entire response
            skills_json = json.loads(skills_text)
            
        # Flatten the skills for backward compatibility
        all_skills = []
        for category in skills_json:
            all_skills.extend(skills_json[category])
            
        return skills_json, all_skills
        
    except (json.JSONDecodeError, AttributeError):
        # Fallback to the old method if JSON parsing fails
        skills = skills_text.split(",")
        skills = [skill.strip() for skill in skills]
        # Create a basic structure with all skills in technical category
        skills_json = {
            "Technical Skills": skills,
            "Soft Skills": [],
            "Domain Knowledge": []
        }
        return skills_json, skills

# def analyze_resume(resume_text, skills, skills_by_category=None):
#     """
#     Enhanced resume analysis with matching by category and detailed feedback
#     """
    
#     # Comprehensive Gemini-based analysis
#     model = genai.GenerativeModel(geminimodel)

#     # If skills_by_category is not provided, assume all skills are technical
#     if skills_by_category is None:
#         skills_by_category = {
#             "Technical Skills": skills,
#             "Soft Skills": [],
#             "Domain Knowledge": []
#         }
    
#     # Basic skill matching for all skills
#     # matched_skills = [skill for skill in skills if skill.lower() in resume_text.lower()]
#     # match_percentage = (len(matched_skills) / len(skills)) * 100 if skills else 0
#     # ── AI-DRIVEN SKILL MATCHING ────────────────────────────────────────    

#     match_prompt = f"""
#     You are an expert resume-parser AI. Given a list of REQUIRED_SKILLS and a RESUME_TEXT, identify exactly which skills are present or clearly implied in the resume.  Handle synonyms, related terms, and context—do not just do substring checks.
    
#     REQUIRED_SKILLS: {skills}
    
#     RESUME_TEXT:
#     {resume_text}
    
#     Return JSON ONLY, in this exact shape:
#     {{
#       "matched_skills": [...],    // skills from REQUIRED_SKILLS that are present/implied
#       "unmatched_skills": [...]   // the rest
#     }}
#     """
#     try:
#         match_resp = model.generate_content(match_prompt)
#         match_json = json.loads(match_resp.text.strip())
#         matched_skills   = match_json.get("matched_skills", [])
#         unmatched_skills = match_json.get("unmatched_skills", [])
#     except Exception:
#         # fallback to literal substring
#         matched_skills   = [s for s in skills if s.lower() in resume_text.lower()]
#         unmatched_skills = [s for s in skills if s not in matched_skills]
#     match_percentage = (len(matched_skills) / len(skills)) * 100 if skills else 0

    
#     # # Category analysis
#     # category_analysis = {}
#     # for category, category_skills in skills_by_category.items():
#     #     if not category_skills:
#     #         category_analysis[category] = {
#     #             "matched": [],
#     #             "unmatched": [],
#     #             "match_percentage": 0
#     #         }
#     #         continue
            
#     #     matched = [skill for skill in category_skills if skill.lower() in resume_text.lower()]
#     #     unmatched = [skill for skill in category_skills if skill.lower() not in resume_text.lower()]
#     #     match_pct = (len(matched) / len(category_skills)) * 100 if category_skills else 0
        
#     #     category_analysis[category] = {
#     #         "matched": matched,
#     #         "unmatched": unmatched,
#     #         "match_percentage": match_pct
#     #     }

#     # Build category_analysis from the AI result
#     category_analysis = {}
#     for category, category_skills in skills_by_category.items():
#         matched   = [s for s in category_skills if s in matched_skills]
#         unmatched = [s for s in category_skills if s in unmatched_skills]
#         pct       = (len(matched)/len(category_skills))*100 if category_skills else 0
#         category_analysis[category] = {
#             "matched": matched,
#             "unmatched": unmatched,
#             "match_percentage": pct
#         }
        
    
#     # # Format skills for the prompt
#     # skills_text = ""
#     # for category, category_skills in skills_by_category.items():
#     #     skills_text += f"{category}:\n"
#     #     for skill in category_skills:
#     #         match_status = "✓" if skill.lower() in resume_text.lower() else "✗"
#     #         skills_text += f"- {match_status} {skill}\n"
    
#     # prompt = f"""
#     # Analyze this resume against the required job skills:
    
#     # RESUME:
#     # {resume_text}
    
#     # REQUIRED SKILLS:
#     # {skills_text}
    
#     # Please provide a comprehensive and detailed analysis with:
#     # 1. A brief overall assessment of how well the resume matches the job skills.
#     # 2. A list of specific recommendations (1D array) to improve the match rate.
#     # 3. A list of priority skills (1D array) that should be added or emphasized in the resume.
#     # 4. A list of sections (1D array) in the resume that need improvement.
#     # 5. ATS score (out of 100) based on matching relevance.

#     # 1. Specific recommendations to improve match rate
#     # 2. Priority skills to add or emphasize
#     # 3. Sections that need improvement
#     # 4. ATS Score should be out of 100 based on matching relevency.
        
#     # Format your response as JSON with these flat keys (no nested objects) as 1D array:
#     # "overall_assessment", "recommendations", "priority_skills", "sections_to_improve", "ats_score"
#     # """

#     # Format skills into ✓/✗ using the AI match results
#     skills_text = ""
#     for cat, cat_skills in skills_by_category.items():
#         skills_text += f"{cat}:\n"
#         for skill in cat_skills:
#             mark = "✓" if skill in matched_skills else "✗"
#             skills_text += f"- {mark} {skill}\n"

#     prompt = f"""
#     You are an expert JD‐Based Resume Tuner AI. Analyze the candidate’s resume against the target job’s required skills and output one flat JSON object with these keys in exactly this order:
    
#       1. overall_assessment (string): 1–2 sentences on fit & gaps.  
#       2. ats_score (integer): 0–100 based on keyword coverage, synonyms, placement, and section weight.  
#       3. keyword_density (object): {"matched": int, "missing": int, "total_required": int}.  
#       4. quick_fixes (array[string]): Top 3 bullet edits deliverable in < 5 minutes.  
#       5. priority_skills (array[string]): Top 5 REQUIRED_SKILLS to highlight immediately.  
#       6. missing_skills (array[string]): REQUIRED_SKILLS not mentioned at all.  
#       7. recommendations (array[string]): Up to 5 medium-term resume rewrites—each flagged High/Med/Low impact.  
#       8. sections_to_improve (array[string]): Exact sections to re-order or rewrite (e.g., "Summary," "Projects").  
#       9. formatting_tips (array[string]): Filetype, layout, font-size, and ATS-friendly design suggestions.  
#      10. redacted_items (array[string]): List of any PII scrubbed (full address, phone, personal email).  
    
#     Inputs (do not hallucinate—use only what’s provided):
#       RESUME_TEXT:
#       {resume_text}
    
#       REQUIRED_SKILLS:
#       {skills_text}
    
#     Rules:
#       • PII: Detect and redact street addresses, personal emails, and phone numbers; list them under "redacted_items."  
#       • Synonyms: Treat common variants (e.g., "AWS" ↔ "Amazon Web Services") as matches for scoring.  
#       • Scoring: Base ats_score on matched vs required keywords, their placement (heading vs body), and section weight (Skills > Experience > Education).  
#       • Impact ranking: Label each recommendation as High/Med/Low based on lift in match rate.  
#       • Output: Return only the JSON object—no extra text, no markdown, no nesting.  
#     """

    
#     try:
#         response = model.generate_content(prompt)
#         analysis_text = response.text.strip()

#         # Try to extract JSON        
#         json_match = re.search(r'```json(.*?)```', analysis_text, re.DOTALL)
#         if json_match:
#             detailed_analysis = json.loads(json_match.group(1).strip())
#         else:
#             # If no JSON code block found, try parsing the entire response
#             detailed_analysis = json.loads(analysis_text)        

#     except:
#         # Fallback in case of parsing issues
#         detailed_analysis = {
#             "overall_assessment": ["This is failsafe callback. Something failed. Check code."],
#             "recommendations": ["Ensure your resume highlights relevant skills explicitly"],
#             "priority_skills": [skill for skill in skills if skill not in matched_skills][:3],
#             "sections_to_improve": ["Skills section", "Work experience"],
#             "ats_score": ["NO ANALYSIS AVAILABLE - Failsafe Exception"]
#         }
    
#     # ── strip markdown from every string in the JSON ──────────────────
#     def _strip_md(txt: str) -> str:
#         # remove **, `, _, ~~ common markdown
#         return re.sub(r'(\*\*|`+|__?|~~)', '', txt)

#     def _clean_md(obj):
#         if isinstance(obj, str):
#             return _strip_md(obj)
#         if isinstance(obj, list):
#             return [_clean_md(v) for v in obj]
#         if isinstance(obj, dict):
#             return {k: _clean_md(v) for k, v in obj.items()}
#         return obj
    
#     detailed_analysis = _clean_md(detailed_analysis)

#     # Emotion-based response (for backward compatibility)
#     if match_percentage < 40:
#         emotion = "😢 Needs improvement"
#     elif match_percentage < 70:
#         emotion = "😊 Good potential"
#     else:
#         emotion = "🎉 Excellent match!"
    
#     # return {
#     #     "matched_skills": matched_skills,
#     #     "unmatched_skills": [skill for skill in skills if skill not in matched_skills],
#     #     "match_percentage": match_percentage,
#     #     "emotion": emotion,
#     #     "category_analysis": category_analysis,
#     #     "detailed_analysis": detailed_analysis
#     # }
#     return {
#         "matched_skills": matched_skills,
#         "unmatched_skills": unmatched_skills,
#         "match_percentage": match_percentage,
#         "emotion": emotion,
#         "category_analysis": category_analysis,
#         "detailed_analysis": detailed_analysis
#     }

def analyze_resume(resume_text, job_description, skills, skills_by_category=None):
    """
    Enhanced resume analysis with matching by category and detailed feedback
    """
    
    # Comprehensive Gemini-based analysis
    model = genai.GenerativeModel(geminimodel)    
    
    # If skills_by_category is not provided, assume all skills are technical
    if skills_by_category is None:
        skills_by_category = {
            "Technical Skills": skills,
            "Soft Skills": [],
            "Domain Knowledge": []
        }
    
    # ── AI-DRIVEN SKILL MATCHING ────────────────────────────────────────    
    match_prompt = f"""
    You are an expert resume-parser AI. Given a list of REQUIRED_SKILLS and a RESUME_TEXT, identify exactly which skills are present or clearly implied in the resume.  Handle synonyms, related terms, and context—do not just do substring checks.
    
    REQUIRED_SKILLS: {skills}
    
    RESUME_TEXT:
    {resume_text}
    
    Return JSON ONLY, in this exact shape:
    {{
      "matched_skills": [...],    // skills from REQUIRED_SKILLS that are present/implied
      "unmatched_skills": [...]   // the rest
    }}
    """   
    try:
        response = model.generate_content(match_prompt)
        # print("AI raw response:", response.text)
        raw = response.text.strip()
        # Try to extract JSON        
        m = re.search(r'```json(.*?)```', raw, re.DOTALL)        
        match_json = json.loads(m.group(1).strip() if m else raw)
              
        matched_skills   = match_json.get("matched_skills", [])        
        unmatched_skills = match_json.get("unmatched_skills", [])
    except Exception as e:
        print("⚠️ Skills AI parse failed:", e)  
        matched_skills   = [s for s in skills if s.lower() in resume_text.lower()]
        unmatched_skills = [s for s in skills if s not in matched_skills]
    match_percentage = round((len(matched_skills) / len(skills)) * 100) if skills else 0

    

    # Build category_analysis from the AI result
    category_analysis = {}
    for category, category_skills in skills_by_category.items():
        matched   = [s for s in category_skills if s in matched_skills]
        unmatched = [s for s in category_skills if s in unmatched_skills]
        pct       = round((len(matched)/len(category_skills))*100) if category_skills else 0
        category_analysis[category] = {
            "matched": matched,
            "unmatched": unmatched,
            "match_percentage": pct
        }
    
    # Format skills into ✓/✗ using the AI match results
    skills_text = ""
    for cat, cat_skills in skills_by_category.items():
        skills_text += f"{cat}:\n"
        for skill in cat_skills:
            mark = "✓" if skill in matched_skills else "✗"
            skills_text += f"- {mark} {skill}\n"

    prompt = f"""
    You are an expert JD‐Based Resume Tuner AI. Analyze the candidate’s resume against the target job’s required skills and output one flat JSON object with these keys in exactly this order:
    
      1. overall_assessment (string): 1–2 sentences on fit & gaps.  
      2. ats_score (integer): 0–100 based on keyword coverage, synonyms, placement, and section weight.  
      3. keyword_density (object): {{"matched": int, "missing": int, "total_required": int}}.
      4. quick_fixes (array[string]): Top 3 bullet edits deliverable in < 5 minutes.  
      5. priority_skills (array[string]): Top 5 REQUIRED_SKILLS to highlight immediately.  
      6. missing_skills (array[string]): REQUIRED_SKILLS not mentioned at all.  
      7. recommendations (array[string]): Up to 5 medium-term resume rewrites—each flagged High/Med/Low impact.  
      8. sections_to_improve (array[string]): Exact sections to re-order or rewrite (e.g., "Summary," "Projects") including details of what is missing.  
      9. formatting_tips (array[string]): Filetype, layout, font-size, and ATS-friendly design suggestions.     
      10. action_verbs (array[string]): Identify and guide users on which lines to change to start with strong, dynamic verbs. 
      11. confidence_score (integer): 0-100 indicating reliability of this analysis.  
      12. tone (string): detected tone of the resume (e.g., professional, enthusiastic).  
      13. behavioral_analysis (object): ratings for key traits, e.g. {{"leadership": string, "teamwork": string, "adaptability": string}}.  
      14. assertiveness_level (integer): 0-100 based on level of assertiveness in language.
      15. clarity (integer): 0-100 based on level of clarity in descriptions.
      16. emotional_intelligence (integer): 0-100 based on presence of emotional intelligence cues.        
      17. customization_level (integer): 0-100 based on degree of tailoring to the job.  
      18. quantification_strength (integer): 0-100 based on strength of numeric data usage.  
      19. readability_score (integer): 0-100 based on readability of the resume.  
      20. grammar_accuracy (integer): 0-100 based on grammar accuracy status.  
      21. structure_coherence (integer): 0-100 based on structural coherence status.  
      22. conciseness (integer): 0-100 based on conciseness evaluation.  
      23. achievement_focus (integer): 0-100 based on focus on achievements.  
      24. leadership_emphasis (integer): 0-100 based on emphasis on leadership.  
      25. teamwork_emphasis (integer): 0-100 based on emphasis on teamwork.  
      26. metric_usage (integer): 0-100 based on frequency of metric usage.
      27. behavioral_score (integer): 0-100 based on ratings for key matching traits and behaviour analysis from the user resume.
    
    Inputs (do not hallucinate—use only what’s provided):
      RESUME_TEXT:
      {resume_text}
    
      REQUIRED_SKILLS:
      {skills_text}

    Rules:      
      • Synonyms: Treat common variants (e.g., "AWS" ↔ "Amazon Web Services") as matches for scoring.  
      • Scoring: Base ats_score on matched vs required keywords, their placement (heading vs body), and section weight (Skills > Experience > Education).  
      • Impact ranking: Label each recommendation as High/Med/Low based on lift in match rate.  
      • Output: Return only the JSON object—no extra text, no markdown, no nesting.  
    """

    try:
        response = model.generate_content(prompt)
        analysis_text = response.text.strip()

        # Try to extract JSON        
        json_match = re.search(r'```json(.*?)```', analysis_text, re.DOTALL)
        if json_match:
            detailed_analysis = json.loads(json_match.group(1).strip())
        else:
            # If no JSON code block found, try parsing the entire response
            detailed_analysis = json.loads(analysis_text)        

    except:
        # Fallback in case of parsing issues
        detailed_analysis = {
            "overall_assessment": ["This is failsafe callback. Something failed. Check code."],
            "recommendations": ["Ensure your resume highlights relevant skills explicitly"],
            "priority_skills": [skill for skill in skills if skill not in matched_skills][:3],
            "sections_to_improve": ["Skills section", "Work experience"],
            "ats_score": ["NO ANALYSIS AVAILABLE - Failsafe Exception"]
        }
        
    def _strip_md(txt: str) -> str:        
        return re.sub(r'(\*\*|`+|__?|~~)', '', txt)

    def _clean_md(obj):
        if isinstance(obj, str):
            return _strip_md(obj)
        if isinstance(obj, list):
            return [_clean_md(v) for v in obj]
        if isinstance(obj, dict):
            return {k: _clean_md(v) for k, v in obj.items()}
        return obj
    
    detailed_analysis = _clean_md(detailed_analysis)

    # Emotion-based response (for backward compatibility)
    # Emotion-based response (for backward compatibility)
    # in your Python emotion logic
    if match_percentage < 10:
        emotion = "Critical Gaps"
    elif match_percentage < 20:
        emotion = "Major Gaps"
    elif match_percentage < 30:
        emotion = "Substantial Gaps"
    elif match_percentage < 40:
        emotion = "Moderate Gaps"
    elif match_percentage < 50:
        emotion = "Minor Gaps"
    elif match_percentage < 60:
        emotion = "Fair Match"
    elif match_percentage < 70:
        emotion = "Good Match"
    elif match_percentage < 80:
        emotion = "Strong Match"
    elif match_percentage < 90:
        emotion = "Excellent Match"
    else:
        emotion = "Outstanding Fit"

     
    return {
        "matched_skills": matched_skills,
        "unmatched_skills": unmatched_skills,
        "match_percentage": match_percentage,
        "emotion": emotion,
        "category_analysis": category_analysis,
        "detailed_analysis": detailed_analysis       
    }

def tailor_resume(resume_text, job_description):
    model = genai.GenerativeModel(geminimodel)
    prompt = f"Given this resume:\n{resume_text}\nand this job description:\n{job_description}\n" \
             "Rewrite the resume to highlight relevant skills and experience, focusing on:" \
             "1. Matching job requirements\n" \
             "2. Emphasizing transferable skills\n" \
             "3. Using industry-specific keywords\n" \
             "4. Maintaining professional tone\n" \
             "5. Structure the resume with clear sections for: Contact Information, Professional Summary, Work Experience, Skills, Education.\n" \
             "6. Format each job entry with company, title, dates, and bullet points for achievements.\n" \
             "7. IMPORTANT: Keep the resume concise to fit on ONE page. Do not use asterisk (*) symbols for bullet points - use hyphens (-) instead.\n" \
             "8. Avoid leaving large gaps between sections and keep descriptions brief but impactful.\n" \
             "9. Use periods at the end of achievement statements only if they form complete sentences."

    response = model.generate_content(prompt)
    return response.text.strip()

def json_resume_to_text(resume_json):
    """
    Given resume_json like:
    { "Contact Information": [...],
      "Professional Summary": [...],
      "Education": [...],
      ... }
    produce a single string with each section title + bullets.
    """
    parts = []
    for section, lines in resume_json.items():
        parts.append(section.upper())
        for line in lines:
            parts.append(f"- {line}")
        parts.append("")  # blank line
    return "\n".join(parts)


@app.route('/')
def home():
    return render_template('index.html')

@app.route('/extract_skills', methods=['POST'])
def get_skills():
    data = request.get_json()
    job_description = sanitize_input(data.get('job_description', ''))
    skills_by_category, all_skills = extract_skills(job_description)
    return jsonify({
        "skills": all_skills,
        "skills_by_category": skills_by_category
    })

@app.route('/create_resume', methods=['POST'])
def create_resume():
    data = request.get_json() or {}
    raw = sanitize_input(data.get('responses',''))
    if not raw:
        return jsonify({"error":"No responses provided"}), 400

    # Prompt Gemini to structure this into an ATS-friendly resume JSON
    model = genai.GenerativeModel(geminimodel)
    prompt = f"""
    The user answered these questions in free form:
    {raw}

    Create a JSON object with these keys (order matters):
    "Contact Information", "Professional Summary", "Education", "Skills",
    "Projects", "Achievements", "Certifications"

    Each value is an array of strings (one per line). 
    Use ATS-friendly, one-page resume structure.
    """
    response = model.generate_content(prompt)
    text = response.text.strip()

    # try to pull out JSON block
    import re, json
    m = re.search(r'```json(.*?)```', text, re.DOTALL)
    j = json.loads(m.group(1)) if m else json.loads(text)

    # convert to plain text and PDF
    pdf_text = json_resume_to_text(j)
    pdf = convert_to_pdf_classic(pdf_text, template_style="professional")

    return send_file(
      pdf,
      download_name='My_Resume.pdf',
      as_attachment=True,
      mimetype='application/pdf'
    )


@app.route('/upload_resume', methods=['POST'])
def upload_resume():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    
    try:
        resume_text = extract_text_from_file(file)
        return jsonify({"resume_text": resume_text})
    except Exception as e:
        return jsonify({"error": str(e)}), 400

@app.route('/analyze_resume', methods=['POST'])
def analyze():
    data = request.get_json()
    resume_text = sanitize_input(data.get('resume_text', ''))
    skills = data.get('skills', [])
    skills_by_category = data.get('skills_by_category', None)
    job_description = sanitize_input(data.get('job_description', ''))
    
    analysis_result = analyze_resume(resume_text, job_description, skills, skills_by_category)
    
    # print( json.dumps(analysis_result, indent=2) )

    return jsonify(analysis_result)

@app.route('/tailor_resume', methods=['POST'])
def tailor():
    # Modified to accept form data instead of JSON
    resume_text = sanitize_input(request.form.get('resume_text', ''))
    job_description = sanitize_input(request.form.get('job_description', ''))
    output_format = sanitize_input(request.form.get('output_format', 'pdf'))
    template_style = sanitize_input(request.form.get('template_style', 'professional'))
    
    tailored_resume = tailor_resume(resume_text, job_description)
    
    if output_format == 'docx':
        docx_file = convert_to_docx_template(tailored_resume, template_style)
        return send_file(docx_file, 
                         download_name='tailored_resume.docx', 
                         as_attachment=True, 
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    else:
        pdf_file = convert_to_pdf_classic(tailored_resume, template_style)
        return send_file(pdf_file, 
                         download_name='tailored_resume.pdf', 
                         as_attachment=True, 
                         mimetype='application/pdf')

@app.route('/preview_resume', methods=['POST'])
def preview_resume():
    data = request.get_json()
    resume_text = sanitize_input(data.get('resume_text', ''))
    job_description = sanitize_input(data.get('job_description', ''))
    
    tailored_resume = tailor_resume(resume_text, job_description)
    return jsonify({"tailored_resume": tailored_resume})

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080, debug=True)